#!/usr/bin/env python3
"""Smoke-test the DOCX renderer against repository examples.

Developer test dependencies:
    python3 -m pip install -r scripts/requirements.txt

This test intentionally uses --no-pdf so it is independent of Word,
LibreOffice, or Pages. PDF conversion is environment-dependent; DOCX content
preservation is the renderer contract this script locks down. End users do not
run this; render_protokoll.py bootstraps its own dependencies.
"""

from __future__ import annotations

import shutil
import subprocess
import sys
import tempfile
from zipfile import ZipFile
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    from docx import Document
    from openpyxl import load_workbook
except ImportError:
    sys.stderr.write(
        "smoke_render.py: missing dependency 'python-docx' or 'openpyxl'.\n"
        "Install with: python3 -m pip install -r scripts/requirements.txt\n"
    )
    sys.exit(3)


REPO_ROOT = Path(__file__).resolve().parents[1]
RENDERER = REPO_ROOT / "scripts" / "render_protokoll.py"
W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
QMG_TEMPLATE_EXAMPLES = {
    "references/examples/beispiel-ausgabe-gespraechsnotiz.md",
    "references/examples/beispiel-ausgabe-eba-interview.md",
    "references/examples/beispiel-ausgabe-einfach.md",
    "references/examples/beispiel-ausgabe-lp1-4.md",
    "references/examples/beispiel-ausgabe-lp5.md",
}
TRACKING_XLSX_EXAMPLES = {
    "references/examples/beispiel-ausgabe-bim.md",
}
XLSX_ONLY_EXAMPLES = TRACKING_XLSX_EXAMPLES


def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def read_xlsx_text(path: Path) -> str:
    wb = load_workbook(path, data_only=False)
    parts: list[str] = []
    for sheet_name in ["Deckblatt", "Protokoll", "Doku_Info"]:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            values = [str(cell.value) for cell in row if cell.value is not None]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def xlsx_template_checks(path: Path, example: str) -> list[str]:
    failures: list[str] = []
    wb = load_workbook(path, data_only=False)
    expected_sheets = ["Deckblatt", "Protokoll", "Doku_Info", "Hilfe und Tipps", "intern"]
    if wb.sheetnames != expected_sheets:
        failures.append(f"{example}: XLSX sheet structure changed: {wb.sheetnames}")
    if "Protokoll" not in wb["Protokoll"].tables:
        failures.append(f"{example}: XLSX Protokoll table is missing")
    else:
        ref = wb["Protokoll"].tables["Protokoll"].ref
        if not ref.startswith("A2:H"):
            failures.append(f"{example}: XLSX Protokoll table should include ausblenden column H (got {ref})")
    if wb["Protokoll"]["H2"].value != "ausblenden":
        failures.append(f"{example}: XLSX ausblenden helper header missing")
    if not str(wb["Protokoll"]["H5"].value or "").startswith("=IF(AND((1+B5)<Deckblatt!$A$3"):
        failures.append(f"{example}: XLSX ausblenden helper formula missing")
    if not isinstance(wb["Deckblatt"]["A3"].value, int):
        failures.append(f"{example}: XLSX meeting number should be numeric in Deckblatt!A3")
    for sheet_name in ["Deckblatt", "Protokoll", "Doku_Info"]:
        sheet_text = "\n".join(
            str(cell.value)
            for row in wb[sheet_name].iter_rows()
            for cell in row
            if cell.value is not None
        )
        for placeholder in [
            "_Vorname_",
            "_Name_",
            "_Firma_",
            "_ Dokument/e, Plan/Pläne _",
            "_Thema 01_",
            "_Ersteller eintragen_",
            "Besprechnungsthema A",
        ]:
            if placeholder in sheet_text:
                failures.append(f"{example}: XLSX placeholder leaked into {sheet_name}: {placeholder}")
    return failures


def simple_xlsx_template_checks(path: Path, example: str) -> list[str]:
    failures: list[str] = []
    wb = load_workbook(path, data_only=False)
    expected_sheets = ["Deckblatt", "Protokoll", "Doku_Info", "Hilfe und Tipps", "intern"]
    if wb.sheetnames != expected_sheets:
        failures.append(f"{example}: simple XLSX sheet structure changed: {wb.sheetnames}")
    if wb["Protokoll"]["A1"].value != "Gesprächsinhalt":
        failures.append(f"{example}: simple XLSX Protokoll sheet is not the official simple template")
    if wb["Protokoll"]["D1"].value != "zuständig" or wb["Protokoll"]["E1"].value != "Frist":
        failures.append(f"{example}: simple XLSX zuständig/Frist headers missing")
    if wb["Protokoll"]["A2"].value != "1" or wb["Protokoll"]["B3"].value != "Thema 01.1":
        failures.append(f"{example}: simple XLSX topic numbering not populated")
    if "Protokoll" in wb["Protokoll"].tables:
        failures.append(f"{example}: simple XLSX should not use the D/K tracking table")
    if wb["Protokoll"]["H2"].value == "ausblenden":
        failures.append(f"{example}: simple XLSX should not include tracking helper column H")
    for sheet_name in ["Deckblatt", "Protokoll", "Doku_Info"]:
        sheet_text = "\n".join(
            str(cell.value)
            for row in wb[sheet_name].iter_rows()
            for cell in row
            if cell.value is not None
        )
        for placeholder in [
            "_Vorname_",
            "_Name_",
            "_Firma_",
            "_Kürzel_",
            "_Prj.-Nr._",
            "_Prj.-Name_",
            "_Besprechungsthema_",
            "_Ersteller_",
            "_Ersteller eintragen_",
            "_ Dokument/e, Plan/Pläne _",
            "_Thema 01_",
            "_Dokument/e, Plan/Pläne_",
            "_Format_",
            "Besprechnungsthema A",
        ]:
            if placeholder in sheet_text:
                failures.append(f"{example}: simple XLSX placeholder leaked into {sheet_name}: {placeholder}")
    return failures


def qmg_template_checks(path: Path, example: str) -> list[str]:
    if example not in QMG_TEMPLATE_EXAMPLES:
        return []
    failures: list[str] = []
    with ZipFile(path) as z:
        names = z.namelist()
        body = ET.fromstring(z.read("word/document.xml")).find("w:body", W_NS)
        if body is None:
            return [f"{example}: DOCX has no body"]
        body_tables = body.findall("w:tbl", W_NS)
        if len(body_tables) != 4:
            failures.append(f"{example}: rendered output should keep exactly 4 QMG body tables")
        xml_blob = "\n".join(
            z.read(name).decode("utf-8", "ignore")
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        )
        for internal_marker in ["Hilfe und Tipps", "Dokument-Raster", "Diese Zeile bitte nicht löschen"]:
            if internal_marker in xml_blob:
                failures.append(f"{example}: internal QMG helper page leaked into output")
        for placeholder in [
            "_Vorname_",
            "_Name_",
            "_Firma_",
            "_ Dokument/e, Plan/Pläne _",
            "_Thema 01_",
            "_Beschreibung einfügen_",
            "Besprechnungsthema A",
        ]:
            if placeholder in xml_blob:
                failures.append(f"{example}: QMG placeholder leaked into output: {placeholder}")
        headers = [name for name in names if name.startswith("word/header")]
        footers = [name for name in names if name.startswith("word/footer")]
        if not headers:
            failures.append(f"{example}: official header parts are missing")
        if not footers:
            failures.append(f"{example}: official footer parts are missing")
        header_text = "\n".join(z.read(name).decode("utf-8", "ignore") for name in headers)
        footer_text = "\n".join(z.read(name).decode("utf-8", "ignore") for name in footers)
        if "Eike Becker_Architekten" not in header_text and "Eike Becker_Architekten" not in footer_text:
            failures.append(f"{example}: EBA header/footer branding missing")
        if "PAGE" not in footer_text or "SECTIONPAGES" not in footer_text:
            failures.append(f"{example}: page number fields missing from footer")
    return failures


def render_example(example: str, required_text: list[str]) -> list[str]:
    failures: list[str] = []
    src = REPO_ROOT / example
    with tempfile.TemporaryDirectory(prefix="eba-render-smoke-") as tmp:
        tmp_path = Path(tmp)
        md_path = tmp_path / src.name
        out_dir = tmp_path / "out"
        shutil.copy2(src, md_path)
        result = subprocess.run(
            [
                sys.executable,
                str(RENDERER),
                str(md_path),
                "--out-dir",
                str(out_dir),
                "--no-pdf",
            ],
            cwd=str(REPO_ROOT),
            text=True,
            capture_output=True,
            timeout=120,
        )
        if result.returncode != 0:
            failures.append(
                f"{example}: renderer exited {result.returncode}: {result.stderr.strip()}"
            )
            return failures
        if md_path.exists():
            failures.append(f"{example}: markdown intermediate was not deleted")

        docx_path = out_dir / f"{md_path.stem}.docx"
        xlsx_path = out_dir / f"{md_path.stem}.xlsx"
        if example in XLSX_ONLY_EXAMPLES:
            if docx_path.exists():
                failures.append(f"{example}: Excel-origin format unexpectedly wrote DOCX")
            if not xlsx_path.exists():
                failures.append(f"{example}: XLSX was not written")
                return failures
            rendered_xlsx = read_xlsx_text(xlsx_path)
            for needle in [
                "BIM-Koordination JF-07",
                "BIMcollab-Issue-Liste",
                "FusionLive bleibt die verbindliche CDE",
                "erledigt",
            ]:
                if needle not in rendered_xlsx:
                    failures.append(f"{example}: rendered XLSX missing {needle!r}")
            failures.extend(xlsx_template_checks(xlsx_path, example))
            return failures

        if not docx_path.exists():
            failures.append(f"{example}: DOCX was not written")
            return failures
        rendered = read_docx_text(docx_path)
        for needle in required_text:
            if needle not in rendered:
                failures.append(f"{example}: rendered DOCX missing {needle!r}")
        failures.extend(qmg_template_checks(docx_path, example))

        if xlsx_path.exists():
            failures.append(f"{example}: unexpected XLSX was written for non-Excel format")
    return failures


def forced_excel_formats() -> list[str]:
    failures: list[str] = []
    cases = [
        (
            "references/examples/beispiel-ausgabe-einfach.md",
            "protokoll-einfach-excel",
            simple_xlsx_template_checks,
            [
                "Kick-Off Meeting Projekt VTS-549",
                "Projektorganisation: EBA übernimmt die Gesamtkoordination.",
                "22.04.26",
            ],
        ),
        (
            "references/examples/beispiel-ausgabe-lp1-4.md",
            "protokoll-lp1-4-excel",
            xlsx_template_checks,
            [
                "Planungsbesprechung — BIM, Bauantrag, Wohnfassade",
                "LP3-Modell (FusionLive-Upload)",
                "DGNB-Workshop am 14.04.26",
            ],
        ),
    ]
    for example, forced_format, check_fn, required_text in cases:
        src = REPO_ROOT / example
        with tempfile.TemporaryDirectory(prefix="eba-render-forced-xlsx-") as tmp:
            tmp_path = Path(tmp)
            md_path = tmp_path / src.name
            out_dir = tmp_path / "out"
            shutil.copy2(src, md_path)
            result = subprocess.run(
                [
                    sys.executable,
                    str(RENDERER),
                    str(md_path),
                    "--format",
                    forced_format,
                    "--out-dir",
                    str(out_dir),
                    "--no-pdf",
                ],
                cwd=str(REPO_ROOT),
                text=True,
                capture_output=True,
                timeout=120,
            )
            if result.returncode != 0:
                failures.append(
                    f"{example} forced {forced_format}: renderer exited "
                    f"{result.returncode}: {result.stderr.strip()}"
                )
                continue
            docx_path = out_dir / f"{md_path.stem}.docx"
            xlsx_path = out_dir / f"{md_path.stem}.xlsx"
            if docx_path.exists():
                failures.append(f"{example} forced {forced_format}: unexpected DOCX was written")
            if not xlsx_path.exists():
                failures.append(f"{example} forced {forced_format}: XLSX was not written")
                continue
            rendered_xlsx = read_xlsx_text(xlsx_path)
            for needle in required_text:
                if needle not in rendered_xlsx:
                    failures.append(
                        f"{example} forced {forced_format}: rendered XLSX missing {needle!r}"
                    )
            failures.extend(check_fn(xlsx_path, f"{example} forced {forced_format}"))
    return failures


def unknown_format_rejected() -> list[str]:
    failures: list[str] = []
    with tempfile.TemporaryDirectory(prefix="eba-render-unknown-") as tmp:
        tmp_path = Path(tmp)
        md_path = tmp_path / "unknown.md"
        out_dir = tmp_path / "out"
        md_path.write_text("# Freitext\n\nDies ist kein EBA-Protokoll.\n", encoding="utf-8")
        result = subprocess.run(
            [
                sys.executable,
                str(RENDERER),
                str(md_path),
                "--out-dir",
                str(out_dir),
                "--no-pdf",
            ],
            cwd=str(REPO_ROOT),
            text=True,
            capture_output=True,
            timeout=120,
        )
        if result.returncode != 3:
            failures.append(f"unknown format: expected renderer exit 3, got {result.returncode}")
        if "Refusing to render without a supported QMG template" not in result.stderr:
            failures.append("unknown format: renderer did not explain QMG-template refusal")
        if (out_dir / "unknown.docx").exists():
            failures.append("unknown format: generic DOCX was written")
    return failures


def main() -> int:
    checks = {
        "references/examples/beispiel-ausgabe-gespraechsnotiz.md": [
            "Gesprächsnotiz",
            "Bauantragsstand und Rückmeldung der Bauaufsicht",
            "Werden innerhalb von 3 Kalendertagen",
        ],
        "references/examples/beispiel-ausgabe-einfach.md": [
            "Protokoll",
            "Kick-Off Meeting Projekt VTS-549",
            "Zuständig/Frist",
        ],
        "references/examples/beispiel-ausgabe-lp1-4.md": [
            "zur Besprechung Nr. 12",
            "Planungsbesprechung — BIM, Bauantrag, Wohnfassade",
            "LP3-Modell (FusionLive-Upload)",
            "Ort | Online",
            "Datum | 24.03.26",
            "Zeit | 09:00 – 09:07",
            "Werden innerhalb von 5 Kalendertagen",
            "Besprechungsthemen",
        ],
        "references/examples/beispiel-ausgabe-lp5.md": [
            "zur Besprechung Nr. 8",
            "Baubesprechung — Rohbau, Mängel, Brandschutz",
            "Witterung",
            "Schalungsplan UG1, UG2",
            "M-048",
        ],
        "references/examples/beispiel-ausgabe-bim.md": [
            "zur Besprechung Nr. 07",
            "BIM-Koordination JF-07",
            "BIMcollab-Issue-Liste",
            "Ort | Online",
            "D/K | B | LN",
        ],
    }

    failures: list[str] = []
    for example, required_text in checks.items():
        failures.extend(render_example(example, required_text))
    failures.extend(forced_excel_formats())
    failures.extend(unknown_format_rejected())

    if failures:
        print(f"Render smoke test failed with {len(failures)} issue(s):")
        for failure in failures:
            print(f"- {failure}")
        return 1
    print("Render smoke test passed.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
